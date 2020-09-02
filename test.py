import os
import shutil
from docx import Document
import win32com
from win32com.client import Dispatch
import re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import math
from scipy import interpolate

seper = os.sep
chord = 0.2       #set chord
wake_num = 70   #wake survey points numbers
surface_pressure_type = 0
wake_static_pressure_type = 1
wake_pressure_type = 2

def ConvertDoc(doc_path, doc_files):
    for doc_file in doc_files:
        if ".doc" in doc_file and ".docx" not in doc_file:
            doc_name = doc_file.strip('.doc')
            word = win32com.client.Dispatch('word.application')
            doc = word.Documents.Open(doc_path + seper + doc_file)
            doc.SaveAs(doc_path + seper + doc_name + '.docx', 12)
            doc.Close()
            os.remove(doc_path + seper + doc_file)

def GenDirName(doc_file):
    pattern = r".+V(?P<velocity>.+)r(?P<AOA>.+)a(?P<BorF>[12]).+"
    pattern = re.compile(pattern)

    three_element = pattern.match(doc_file, re.IGNORECASE)
    if three_element.group(3) == '1':
        dir_name = 'V' + three_element.group(1) + 'A' + three_element.group(2) + 'Forward'
    elif three_element.group(3) == '2':
        dir_name = 'V' + three_element.group(1) + 'A' + three_element.group(2) + 'Backward'

    return dir_name

def GenDirs(doc_path, doc_files):
    for doc_file in doc_files:
        if ".docx" in doc_file and "error" not in doc_file:
            dir_name = GenDirName(doc_file)
            
            if dir_name not in os.listdir(doc_path):
                os.mkdir(doc_path + seper + dir_name)
        
def OutputPressure(doc_path, doc_files):
    p_oo_err, p_o_err, pressure_err = GetErrorPressure(doc_path, doc_files, surface_pressure_type)

    for doc_file in doc_files:
        coordinate = []
        pressure = []
        
        if ".docx" in doc_file and "error" not in doc_file:
            document = Document(doc_path + seper + doc_file)
            tables = [table for table in document.tables]
            print(len(tables))
            print(doc_file)
            p_oo, p_o = float(tables[3].rows[2].cells[4].text) - p_oo_err, float(tables[3].rows[1].cells[4].text) - p_o_err
            
            for row in tables[0].rows[1:]:
                coordinate.append(row.cells[1].text)
                pressure.append(row.cells[4].text)

            coordinate = list(map(float, coordinate))
            pressure = list(map(float, pressure))
            pressure = list(map(lambda p: p[0] - p[1], zip(pressure, pressure_err)))

            C_p = [(p - p_oo) / (p_o - p_oo) for p in pressure]

            pres_df = pd.DataFrame({'coordinate':coordinate, 'pressure':pressure, 'Cp':C_p})
            dir_name = GenDirName(doc_file)
            pres_df.to_csv(doc_path + seper + dir_name + seper + dir_name + '.csv', sep = '\t', index=False)

def GetErrorPressure(doc_path, doc_files, pressuretype):
	pressure_err = []

	if "error.docx" in doc_files:
		document = Document(doc_path + seper + "error.docx")
		tables = [table for table in document.tables]
		p_oo_err, p_o_err = float(tables[3].rows[2].cells[4].text), float(tables[3].rows[1].cells[4].text)

		for row in tables[pressuretype].rows[1:]:
			pressure_err.append(row.cells[4].text)

		pressure_err = list(map(float, pressure_err))

	return p_oo_err, p_o_err, pressure_err

def C_d(y, p_i, p_oi, p_oo, p_o):
    c_x = [math.sqrt(math.fabs((p_oi[i] - p_i[i]) / (p_o - p_oo))) * (1 - math.sqrt(math.fabs((p_oi[i] - p_i[i]) / (p_o - p_oo)))) for i in range(wake_num + 1)]
    c_d = np.trapz(c_x, y) * 2 / chord
    return c_d

def GetAllCLCD(doc_path, doc_files):
    p_oo_err, p_o_err, static_pressure_err = GetErrorPressure(doc_path, doc_files, wake_static_pressure_type)
    none1, none2, pressure_err = GetErrorPressure(doc_path, doc_files, wake_pressure_type)
    del none1
    del none2

    AOA_F = []
    AOA_B = []
    C_D_F = []
    C_L_F = []
    C_D_B = []
    C_L_B = []
    LDratio_F = []
    LDratio_B = []

    for doc_file in doc_files:
        static_y = []
        y = []
        static_pressure = []
        pressure = []
        
        if ".docx" in doc_file and "error" not in doc_file:
            document = Document(doc_path + seper + doc_file)
            tables = [table for table in document.tables]
            p_oo, p_o = float(tables[3].rows[2].cells[4].text) - p_oo_err, float(tables[3].rows[1].cells[4].text) - p_o_err
    
            for row in tables[wake_static_pressure_type].rows[1:]:
                static_y.append(row.cells[1].text)
                static_pressure.append(row.cells[4].text)

            static_y = list(map(float, static_y))
            static_y = [x * chord for x in static_y]
            static_pressure = list(map(float, static_pressure))
            static_pressure = list(map(lambda p: p[0] - p[1], zip(static_pressure, static_pressure_err)))

            for row in tables[wake_pressure_type].rows[1:]:
                y.append(row.cells[1].text)
                pressure.append(row.cells[4].text)

            y = list(map(float, y))
            y = [x * chord for x in y]
            pressure = list(map(float, pressure))
            pressure = list(map(lambda p: p[0] - p[1], zip(pressure, pressure_err)))

            tck = interpolate.splrep(static_y, static_pressure) #simulation
            static_pressure = interpolate.splev(y, tck)

            C_D = C_d(y, static_pressure, pressure, p_oo, p_o)
            
            pattern = r".+V(?P<velocity>.+)r(?P<AOA>.+)a(?P<BorF>[12]).+"
            pattern = re.compile(pattern)
            three_element = pattern.match(doc_file, re.IGNORECASE)

            alpha = math.radians(float(three_element.group(2)))   #degree to rad

            dir_name = GenDirName(doc_file)
            cp_df = pd.read_csv(doc_path + seper + dir_name + seper + dir_name + '.csv', sep = '\t', usecols=['coordinate', 'Cp'])
            x = cp_df['coordinate'] * chord
            c_p = cp_df['Cp']
            C_n = np.trapz(x, c_p)

            C_L = C_n * math.cos(alpha) / chord - (C_D - C_n * math.sin(alpha) / chord) * math.tan(alpha)

            LDratio = C_L / C_D

            if three_element.group(3) == '1':
                AOA_F.append(float(three_element.group(2)))
                C_D_F.append(C_D)
                C_L_F.append(C_L)
                LDratio_F.append(LDratio)
            elif three_element.group(3) == '2':
                AOA_B.append(float(three_element.group(2)))
                C_D_B.append(C_D)
                C_L_B.append(C_L)
                LDratio_B.append(LDratio)

    #forward_df = pd.DataFrame({'AOA': AOA_F, 'Cd': C_D_F, 'Cl': C_L_F, 'L/D': LDratio_F}).sort_values(by='AOA')
    backward_df = pd.DataFrame({'AOA': AOA_B, 'Cd': C_D_B, 'Cl': C_L_B, 'L/D': LDratio_B}).sort_values(by='AOA')
    
    ax = plt.gca()
    #forward_df.plot(kind='line', x='AOA', y='Cl', label='forward', ax=ax)
    backward_df.plot(kind='line', x='AOA', y='Cl', label='backward', color='orange', ax=ax)
    plt.show()

    ax = plt.gca()
    #forward_df.plot(kind='line', x='AOA', y='Cd', label='forward', ax=ax)
    backward_df.plot(kind='line', x='AOA', y='Cd', label='backward', color='orange', ax=ax)
    plt.show()

    ax = plt.gca()
    #forward_df.plot(kind='line', x='AOA', y='L/D', label='forward', ax=ax)
    backward_df.plot(kind='line', x='AOA', y='L/D', label='backward', color='orange', ax=ax)
    plt.show()

    #forward_df.to_csv(doc_path + seper + 'forward.csv', sep = ',', index=False)
    backward_df.to_csv(doc_path + seper + 'backward.csv', sep = ',', index=False)
                
if __name__ == "__main__":
    doc_path = os.getcwd()
    doc_files = os.listdir(doc_path)
    ConvertDoc(doc_path, doc_files)
    GenDirs(doc_path, doc_files)
    OutputPressure(doc_path, doc_files)
    GetAllCLCD(doc_path, doc_files)
    
